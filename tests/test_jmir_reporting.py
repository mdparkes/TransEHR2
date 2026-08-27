"""Tests for the JMIR manuscript reporting utilities.

Covers the three parts that a wrong answer would silently corrupt: the
JMIR number formatting rules, the corrected resampled t test and the
Benjamini-Hochberg adjustment, and the metric computations. The
concordance index is checked against the O(n^2) pair-enumeration version
in ``evaluate_finetuned_predictions.py``, which is the definition of
record.
"""

import math
import os
import sys

import numpy as np
import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from evaluate_finetuned_predictions import concordance_index as brute_force
from jmir_reporting.evaluation import (
    binary_metrics,
    calibrate_thresholds,
    concordance_index,
    phenotype_metrics,
    prevalence_matched_threshold,
)
from jmir_reporting.formatting import (
    fmt_cell,
    fmt_number,
    fmt_p_value,
    fmt_t_statistic,
)
from jmir_reporting.stats import (
    benjamini_hochberg,
    corrected_resampled_ttest,
    standard_error_of_mean,
)
from jmir_reporting.tables import (
    Table,
    build_document,
    parse_markup,
    render_text,
    strip_markup,
)


# ------------------------------------------------------------------
# JMIR formatting rules
# ------------------------------------------------------------------

@pytest.mark.parametrize('value,expected', [
    (0.03, 'P=.03'),
    (0.0271, 'P=.03'),
    (0.5, 'P=.50'),
    (0.08094, 'P=.08'),
    # P<.01 takes a third decimal place.
    (0.003, 'P=.003'),
    (0.0086, 'P=.009'),
    # Rounding to 2 places would cross alpha, so keep 3.
    (0.048, 'P=.048'),
    (0.0475, 'P=.048'),
    # A P value can be neither 0 nor 1.
    (0.0000115, 'P<.001'),
    (0.0, 'P<.001'),
    (1.0, 'P>.99'),
    (0.995, 'P>.99'),
])
def test_p_values_follow_jmir_rules(value, expected):
    assert fmt_p_value(value) == expected


def test_p_value_has_no_leading_zero():
    assert not fmt_p_value(0.03).startswith('P=0')


def test_p_value_unavailable_renders_as_dash():
    assert fmt_p_value(float('nan')) == '—'
    assert fmt_p_value(None) == '—'


def test_numbers_below_one_keep_their_leading_zero():
    assert fmt_number(0.836) == '0.836'


def test_negative_numbers_use_an_en_dash():
    assert fmt_number(-0.043, 3) == '–0.043'
    assert '-' not in fmt_number(-0.043, 3)


def test_cell_separates_unrelated_statistics_with_a_semicolon():
    assert fmt_cell(0.845, 0.004, 0.0271) == '0.845 (SE 0.004; P=.03)'


def test_control_cell_carries_no_p_value():
    assert fmt_cell(0.858, 0.002, None) == '0.858 (SE 0.002)'


def test_t_statistic_puts_df_in_parentheses():
    assert fmt_t_statistic(2.6789, 15) == '2.68 (15)'


# ------------------------------------------------------------------
# Corrected resampled t test
# ------------------------------------------------------------------

def test_correction_inflates_the_standard_error_as_theory_predicts():
    """At k folds the statistic shrinks by sqrt((1/k + 1/(k-1)) / (1/k))."""
    rng = np.random.default_rng(0)
    for k in (3, 5, 10):
        a = rng.normal(0.85, 0.01, k)
        b = rng.normal(0.83, 0.01, k)
        result = corrected_resampled_ttest(a, b)
        d = a - b
        paired = d.mean() / (d.std(ddof=1) / math.sqrt(k))
        expected = math.sqrt((1 / k + 1 / (k - 1)) / (1 / k))
        assert result.df == k - 1
        assert paired / result.statistic == pytest.approx(expected, abs=1e-12)


def test_fixed_adjustment_is_one_over_k_minus_one():
    """The default matches an explicit n_test/n_train of 1/(k-1)."""
    rng = np.random.default_rng(1)
    a, b = rng.normal(size=5), rng.normal(size=5)
    default = corrected_resampled_ttest(a, b)
    explicit = corrected_resampled_ttest(a, b, n_train_test_ratio=1 / 4)
    assert default.statistic == pytest.approx(explicit.statistic)


def test_the_test_is_paired_not_unpaired():
    """Reordering one arm's folds must change the result."""
    a = [0.86, 0.85, 0.85, 0.84, 0.84]
    b = [0.84, 0.83, 0.85, 0.82, 0.83]
    original = corrected_resampled_ttest(a, b).statistic
    shuffled = corrected_resampled_ttest(a, list(reversed(b))).statistic
    assert original != pytest.approx(shuffled)


def test_identical_arms_give_no_p_value():
    """Prevalence is identical in every column; P can never equal 1."""
    result = corrected_resampled_ttest([0.17] * 5, [0.17] * 5)
    assert math.isnan(result.p_value)
    assert result.note


def test_constant_non_zero_difference_is_reported_as_highly_significant():
    result = corrected_resampled_ttest([0.2] * 5, [0.1] * 5)
    assert result.p_value == 0.0
    assert fmt_p_value(result.p_value) == 'P<.001'


def test_nan_folds_are_dropped_and_reduce_the_df():
    result = corrected_resampled_ttest(
        [0.86, float('nan'), 0.85, 0.84, 0.84],
        [0.84, 0.83, 0.85, 0.82, 0.83]
    )
    assert result.df == 3


def test_mismatched_fold_counts_are_rejected():
    with pytest.raises(ValueError):
        corrected_resampled_ttest([0.1, 0.2, 0.3], [0.1, 0.2])


def test_too_few_folds_are_rejected():
    with pytest.raises(ValueError):
        corrected_resampled_ttest([0.1], [0.2])


# ------------------------------------------------------------------
# Benjamini-Hochberg
# ------------------------------------------------------------------

def test_bh_matches_r_p_adjust():
    """Values verified against R's p.adjust(method = "BH")."""
    assert benjamini_hochberg([0.01, 0.02, 0.03, 0.04, 0.05]) == \
        pytest.approx([0.05] * 5)
    assert benjamini_hochberg([0.001, 0.008, 0.039, 0.041, 0.042]) == \
        pytest.approx([0.005, 0.02, 0.042, 0.042, 0.042])
    # Order of the input must not matter.
    assert benjamini_hochberg([0.042, 0.001, 0.041, 0.008, 0.039]) == \
        pytest.approx([0.042, 0.005, 0.042, 0.02, 0.042])


def test_bh_is_monotone():
    rng = np.random.default_rng(3)
    raw = rng.random(40)
    adjusted = benjamini_hochberg(raw)
    order = np.argsort(raw)
    assert np.all(np.diff(adjusted[order]) >= -1e-12)


def test_bh_never_exceeds_one():
    assert np.all(benjamini_hochberg([0.9, 0.95, 0.99]) <= 1.0)


def test_bh_passes_nan_through_and_excludes_it_from_the_family():
    adjusted = benjamini_hochberg([0.01, 0.02, float('nan')])
    assert math.isnan(adjusted[2])
    # Two tests, not three: 0.01 * 2 / 1 = 0.02.
    assert adjusted[0] == pytest.approx(0.02)


def test_sem_uses_the_sample_standard_deviation():
    values = [0.86, 0.85, 0.85, 0.84, 0.84]
    expected = np.std(values, ddof=1) / math.sqrt(5)
    assert standard_error_of_mean(values) == pytest.approx(expected)


# ------------------------------------------------------------------
# Concordance index
# ------------------------------------------------------------------

@pytest.mark.parametrize('seed', range(6))
def test_concordance_index_matches_pair_enumeration(seed):
    rng = np.random.default_rng(seed)
    n = int(rng.integers(10, 70))
    observed = rng.integers(0, 8, n).astype(float)   # ties in observed
    predicted = np.round(rng.normal(0, 1, n), 1)     # ties in predicted
    assert concordance_index(predicted, observed) == \
        pytest.approx(brute_force(predicted, observed), abs=1e-12)


def test_concordance_index_of_a_perfect_ranking_is_one():
    observed = np.arange(50, dtype=float)
    assert concordance_index(observed, observed) == pytest.approx(1.0)


def test_concordance_index_of_a_reversed_ranking_is_zero():
    observed = np.arange(50, dtype=float)
    assert concordance_index(-observed, observed) == pytest.approx(0.0)


def test_concordance_index_needs_comparable_pairs():
    assert math.isnan(concordance_index([0.1, 0.2], [5.0, 5.0]))
    assert math.isnan(concordance_index([0.1], [5.0]))


# ------------------------------------------------------------------
# Threshold calibration
# ------------------------------------------------------------------

def test_prevalence_matching_reproduces_the_observed_positive_rate():
    rng = np.random.default_rng(4)
    probabilities = rng.random(2000)
    targets = (rng.random(2000) < 0.17).astype(float)
    threshold = prevalence_matched_threshold(probabilities, targets)
    assert (probabilities >= threshold).mean() == pytest.approx(
        targets.mean()
    )


def test_prevalence_matching_falls_back_when_a_label_never_occurs():
    assert prevalence_matched_threshold([0.1, 0.2, 0.3], [0, 0, 0]) == 0.5
    assert prevalence_matched_threshold([0.1, 0.2, 0.3], [1, 1, 1]) == 0.5


def test_per_label_calibration_gives_one_threshold_per_label():
    rng = np.random.default_rng(5)
    probabilities = rng.random((500, 4))
    targets = (rng.random((500, 4)) < [0.4, 0.2, 0.1, 0.02]).astype(float)
    thresholds = calibrate_thresholds(
        probabilities, targets, 'prevalence', 'per-label'
    )
    assert thresholds.shape == (4,)
    # A rarer label needs a higher threshold to keep its positive count.
    assert thresholds[3] > thresholds[0]


def test_global_calibration_shares_one_threshold():
    rng = np.random.default_rng(6)
    probabilities = rng.random((500, 4))
    targets = (rng.random((500, 4)) < 0.2).astype(float)
    thresholds = calibrate_thresholds(
        probabilities, targets, 'prevalence', 'global'
    )
    assert np.all(thresholds == thresholds[0])


def test_a_fixed_threshold_is_used_verbatim():
    thresholds = calibrate_thresholds(
        np.zeros((10, 3)), np.zeros((10, 3)), 0.5, 'per-label'
    )
    assert np.all(thresholds == 0.5)


def test_an_unknown_threshold_mode_is_rejected():
    with pytest.raises(ValueError):
        calibrate_thresholds(np.zeros((4, 1)), np.zeros((4, 1)), 'youden')


# ------------------------------------------------------------------
# Metrics
# ------------------------------------------------------------------

def test_binary_metrics_are_internally_consistent():
    rng = np.random.default_rng(7)
    y = (rng.random(400) < 0.3).astype(float)
    p = rng.random(400)
    m = binary_metrics(y, p, 0.4)
    assert m['false_positive_rate'] == pytest.approx(1 - m['specificity'])
    assert m['false_negative_rate'] == pytest.approx(
        1 - m['recall_sensitivity']
    )
    assert m['false_discovery_rate'] == pytest.approx(1 - m['ppv'])
    assert m['ppv'] == pytest.approx(m['precision'])
    assert m['prevalence'] == pytest.approx(y.mean())


def test_auroc_is_unaffected_by_the_threshold():
    rng = np.random.default_rng(8)
    y = (rng.random(300) < 0.3).astype(float)
    p = rng.random(300)
    assert binary_metrics(y, p, 0.2)['auroc'] == \
        pytest.approx(binary_metrics(y, p, 0.8)['auroc'])


def test_phenotype_accuracy_is_the_same_micro_and_macro():
    """Every label contributes the same number of predictions."""
    rng = np.random.default_rng(9)
    probabilities = rng.random((200, 6))
    targets = (rng.random((200, 6)) < 0.25).astype(float)
    m = phenotype_metrics(probabilities, targets, np.full(6, 0.5))
    assert m['micro_accuracy'] == pytest.approx(m['macro_accuracy'])


def test_phenotype_honours_a_different_threshold_per_label():
    probabilities = np.tile(np.array([[0.4, 0.4]]), (10, 1))
    targets = np.zeros((10, 2))
    lenient = phenotype_metrics(probabilities, targets, np.array([0.3, 0.9]))
    # Only the first label crosses its threshold, so half the pooled
    # predictions are positive.
    assert lenient['micro_false_positive_rate'] == pytest.approx(0.5)


# ------------------------------------------------------------------
# Agreement with evaluate_finetuned_predictions.py
# ------------------------------------------------------------------

def _same_values(left, right):
    """Compare two metric dicts, treating NaN as equal to NaN.

    Args:
        left: A metric dict.
        right: Another metric dict with the same keys.

    Returns:
        A list of the keys whose values disagree.
    """
    assert set(left) == set(right), set(left) ^ set(right)
    return [
        key for key in left
        if not (math.isnan(left[key]) and math.isnan(right[key]))
        and abs(left[key] - right[key]) > 1e-12
    ]


@pytest.mark.parametrize('threshold', [0.3, 0.5, 0.7])
def test_binary_metrics_match_the_evaluation_script(threshold):
    from evaluate_finetuned_predictions import compute_binary_metrics

    rng = np.random.default_rng(11)
    y = (rng.random(800) < 0.17).astype(float)
    p = rng.random(800)
    assert _same_values(
        compute_binary_metrics(y, p, threshold),
        binary_metrics(y, p, threshold)
    ) == []


def test_phenotype_metrics_match_the_evaluation_script():
    """Checked at a shared threshold, the only mode the older script has."""
    from evaluate_finetuned_predictions import compute_phenotype_metrics

    rng = np.random.default_rng(12)
    n_labels = 25
    probabilities = rng.random((600, n_labels))
    targets = (
        rng.random((600, n_labels)) < np.linspace(0.45, 0.01, n_labels)
    ).astype(float)
    assert _same_values(
        compute_phenotype_metrics(probabilities, targets, 0.5),
        phenotype_metrics(probabilities, targets, np.full(n_labels, 0.5))
    ) == []


def test_length_of_stay_metrics_match_the_evaluation_script():
    from evaluate_finetuned_predictions import compute_los_metrics
    from jmir_reporting.evaluation import length_of_stay_metrics

    rng = np.random.default_rng(13)
    observed = np.abs(rng.lognormal(4.4, 0.8, 300)).reshape(-1, 1)
    predicted = np.abs(observed + rng.normal(0, 40, (300, 1)))
    assert _same_values(
        compute_los_metrics(predicted, observed),
        length_of_stay_metrics(predicted, observed)
    ) == []


# ------------------------------------------------------------------
# Table rendering
# ------------------------------------------------------------------

def test_markup_parses_into_formatted_runs():
    assert parse_markup('<i>F</i><sub>1</sub>-score') == [
        ('F', {'i'}), ('1', {'sub'}), ('-score', set())
    ]


def test_stripping_markup_keeps_footnote_markers_legible():
    assert strip_markup('Metric<sup>a,b</sup>') == 'Metric[a,b]'
    assert strip_markup('<i>F</i><sub>1</sub>-score') == 'F1-score'


def test_footnote_letters_are_allocated_in_order():
    table = Table(1, 'Caption', 'Metric', ['A', 'B'])
    assert table.add_footnote('first') == 'a'
    assert table.add_footnote('second') == 'b'


def test_a_table_with_categories_is_nested():
    table = Table(1, 'Caption', 'Metric', ['A'])
    table.add_row('Accuracy', ['0.8'])
    assert not table.is_nested
    table.add_category('Microaverages')
    table.add_row('AUROC', ['0.7'], level=1)
    assert table.is_nested


def test_word_output_nests_categories_and_repeats_the_header(tmp_path):
    from docx import Document

    table = Table(3, 'Diagnosis prediction', 'Metric<sup>a</sup>',
                  ['In-stay records only', 'Historical records only'])
    table.add_footnote(' A note.')
    table.add_row('Accuracy', ['0.810 (SE 0.000)', '0.788 (SE 0.001)'])
    table.add_category('Microaverages')
    table.add_row('<i>F</i><sub>1</sub>-score',
                  ['0.424 (SE 0.004)', '0.243 (SE 0.010)'], level=1)

    path = tmp_path / 'table.docx'
    build_document(table, str(path))
    grid = Document(str(path)).tables[0]

    # Two stub columns plus two data columns.
    assert len(grid.columns) == 4

    def distinct(row):
        seen, out = set(), []
        for cell in row.cells:
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            out.append(cell)
        return out

    header, accuracy, category, f1 = (distinct(r) for r in grid.rows)

    # The stub heading spans both stub columns.
    assert len(header) == 3
    assert header[0].text == 'Metrica'
    # A top-level metric row also spans both stub columns.
    assert len(accuracy) == 3
    # The category heading is bold and merged with the subcategory column.
    assert len(category) == 3
    assert category[0].text == 'Microaverages'
    assert category[0].paragraphs[0].runs[0].bold
    # A subcategory row leaves the first column empty.
    assert len(f1) == 4
    assert f1[0].text == ''
    assert f1[1].text == 'F1-score'
    runs = f1[1].paragraphs[0].runs
    assert runs[0].italic and runs[1].font.subscript

    header_row = grid.rows[0]._tr.find(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr'
    )
    assert header_row is not None and len(header_row) > 0


def test_text_rendering_uses_short_column_labels(capsys):
    table = Table(1, 'Caption', 'Metric',
                  ['A very long column heading indeed', 'Another one'],
                  short_columns=['Expt 3', 'Expt 1'])
    table.add_row('Accuracy', ['0.8', '0.7'])
    render_text(table)
    out = capsys.readouterr().out
    assert 'Expt 3 = A very long column heading indeed' in out
    assert max(len(line) for line in out.splitlines()) < 80
