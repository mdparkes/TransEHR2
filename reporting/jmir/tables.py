"""Build, print and export manuscript tables in JMIR house style.

A table is described once as a :class:`Table` and then rendered either to
stdout for inspection or to a Word document for submission. The Word
output follows JMIR Publications' table guidelines:

* the table is built with the Word table function and preceded by a
  numbered caption,
* column and row headings are in sentence case, with units of
  measurement separated from the description by a comma,
* category headings are bold and merged with the subcategory column to
  preserve indentation,
* footnote markers are superscript letters running left to right and top
  to bottom, never asterisks tied to significance levels, and
* the header row repeats when the table breaks across pages.

Row and cell text may contain the inline markup ``<i>``, ``<b>``,
``<sub>`` and ``<sup>``, which the Word renderer turns into character
formatting and the text renderer strips. This is what produces a correct
*F*₁-score, with an italic F and a subscript 1.
"""

import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

FOOTNOTE_LETTERS = 'abcdefghijklmnopqrstuvwxyz'

_TAG = re.compile(r'</?(i|b|sub|sup)>')


# ------------------------------------------------------------------
# Table model
# ------------------------------------------------------------------

class Row:
    """One row of a table.

    Attributes:
        label: The row heading, which may contain inline markup.
        cells: One string per data column. Empty for category headings.
        kind: ``'metric'`` for a data row or ``'category'`` for a bold
            grouping heading such as "Microaverages".
        level: ``0`` for a top-level row heading and ``1`` for a row that
            sits underneath a category heading.
    """

    __slots__ = ('label', 'cells', 'kind', 'level')

    def __init__(self, label, cells=None, kind='metric', level=0):
        self.label = label
        self.cells = list(cells) if cells else []
        self.kind = kind
        self.level = level


class Table:
    """A complete manuscript table.

    Attributes:
        number: The table number used in the caption.
        caption: The caption text, without the "Table N." prefix.
        stub_head: Heading for the row-heading column.
        columns: One heading per data column.
        short_columns: Optional abbreviated headings used only by
            :func:`render_text`, so that a wide table stays readable in a
            terminal. The Word output always uses ``columns``.
        rows: The :class:`Row` objects, in order.
        footnotes: Footnote texts, in the order their letters were
            allocated.
    """

    def __init__(self, number, caption, stub_head, columns,
                 short_columns=None):
        self.number = number
        self.caption = caption
        self.stub_head = stub_head
        self.columns = list(columns)
        self.short_columns = list(short_columns) if short_columns else None
        self.rows = []
        self.footnotes = []

    def add_footnote(self, text):
        """Register a footnote and return its superscript letter.

        Footnote letters are allocated in call order, so callers should
        register them as they build the table from left to right and top
        to bottom, which is the order JMIR requires.

        Args:
            text: The footnote text, which may contain inline markup.

        Returns:
            The footnote letter, e.g. ``'a'``.

        Raises:
            ValueError: If more than 26 footnotes are requested.
        """
        if len(self.footnotes) >= len(FOOTNOTE_LETTERS):
            raise ValueError('A table may carry at most 26 footnotes')
        self.footnotes.append(text)
        return FOOTNOTE_LETTERS[len(self.footnotes) - 1]

    def add_row(self, label, cells, level=0):
        """Append a data row.

        Args:
            label: The row heading.
            cells: One string per data column.
            level: ``0`` or ``1``; see :class:`Row`.
        """
        self.rows.append(Row(label, cells, 'metric', level))

    def add_category(self, label):
        """Append a bold category heading row.

        Args:
            label: The category heading, e.g. ``'Microaverages'``.
        """
        self.rows.append(Row(label, [], 'category', 0))

    @property
    def is_nested(self):
        """Whether any row sits under a category heading."""
        return any(row.kind == 'category' or row.level > 0
                   for row in self.rows)


# ------------------------------------------------------------------
# Inline markup
# ------------------------------------------------------------------

def strip_markup(text):
    """Remove inline markup, leaving plain text.

    Superscripts become bracketed so that footnote markers stay legible
    in a terminal, where "Metrica" would otherwise read as one word.

    Args:
        text: Text possibly containing ``<i>``, ``<b>``, ``<sub>`` or
            ``<sup>`` tags.

    Returns:
        The text with all tags removed.
    """
    text = re.sub(r'<sup>(.*?)</sup>', r'[\1]', text)
    return _TAG.sub('', text)


def parse_markup(text):
    """Split marked-up text into runs with their character formatting.

    Args:
        text: Text possibly containing ``<i>``, ``<b>``, ``<sub>`` or
            ``<sup>`` tags.

    Returns:
        A list of ``(text, style)`` tuples, where ``style`` is a set
        drawn from ``{'i', 'b', 'sub', 'sup'}``.
    """
    runs = []
    active = []
    position = 0

    for match in _TAG.finditer(text):
        chunk = text[position:match.start()]
        if chunk:
            runs.append((chunk, set(active)))
        tag = match.group(1)
        if match.group(0).startswith('</'):
            if active and active[-1] == tag:
                active.pop()
            elif tag in active:
                active.remove(tag)
        else:
            active.append(tag)
        position = match.end()

    tail = text[position:]
    if tail:
        runs.append((tail, set(active)))
    return runs


def _add_runs(paragraph, text, bold=False):
    """Write marked-up text into a Word paragraph as formatted runs.

    Args:
        paragraph: The ``docx`` paragraph to append to.
        text: Text possibly containing inline markup.
        bold: Apply bold to every run, for category and header rows.
    """
    for chunk, style in parse_markup(text):
        run = paragraph.add_run(chunk)
        run.italic = 'i' in style
        run.bold = bold or ('b' in style)
        if 'sub' in style:
            run.font.subscript = True
        if 'sup' in style:
            run.font.superscript = True


# ------------------------------------------------------------------
# Plain-text rendering
# ------------------------------------------------------------------

def render_text(table, stream=None):
    """Render a table as aligned plain text.

    Args:
        table: The :class:`Table` to render.
        stream: A writable stream; defaults to stdout.

    Returns:
        The rendered text.
    """
    import sys

    if stream is None:
        stream = sys.stdout

    headings = table.short_columns or table.columns
    header = [strip_markup(table.stub_head)]
    header += [strip_markup(c) for c in headings]

    body = []
    for row in table.rows:
        label = strip_markup(row.label)
        if row.level > 0:
            label = '    ' + label
        if row.kind == 'category':
            body.append([label] + [''] * len(table.columns))
        else:
            body.append([label] + [strip_markup(c) for c in row.cells])

    widths = [len(h) for h in header]
    for line in body:
        for i, cell in enumerate(line):
            widths[i] = max(widths[i], len(cell))

    def format_line(cells):
        parts = [cells[0].ljust(widths[0])]
        parts += [cells[i].ljust(widths[i]) for i in range(1, len(cells))]
        return '  '.join(parts).rstrip()

    lines = [f'Table {table.number}. {strip_markup(table.caption)}', '']
    if table.short_columns:
        lines.append('Columns:')
        for short, full in zip(table.short_columns, table.columns):
            lines.append(f'  {strip_markup(short)} = {strip_markup(full)}')
        lines.append('')
    lines.append(format_line(header))
    lines.append('-' * min(len(lines[-1]), 200))
    for line in body:
        lines.append(format_line(line))
    if table.footnotes:
        lines.append('')
        for letter, note in zip(FOOTNOTE_LETTERS, table.footnotes):
            lines.append(f'{letter}) {strip_markup(note)}')

    text = '\n'.join(lines)
    print(text, file=stream)
    return text


# ------------------------------------------------------------------
# Word rendering
# ------------------------------------------------------------------

def _repeat_header_row(row):
    """Mark a table row as a header that repeats across page breaks.

    Args:
        row: The ``docx`` table row.
    """
    properties = row._tr.get_or_add_trPr()
    marker = properties.makeelement(qn('w:tblHeader'), {})
    properties.append(marker)


def _write_cell(cell, text, bold=False, align=None):
    """Write marked-up text into a table cell.

    Args:
        cell: The ``docx`` table cell.
        text: Text possibly containing inline markup.
        bold: Apply bold to every run.
        align: Optional ``WD_ALIGN_PARAGRAPH`` value.
    """
    paragraph = cell.paragraphs[0]
    # Assigning paragraph.text would leave behind one empty, unformatted
    # run, so drop the existing runs from the XML instead.
    for existing in list(paragraph.runs):
        existing._element.getparent().remove(existing._element)
    if align is not None:
        paragraph.alignment = align
    _add_runs(paragraph, text, bold=bold)


def build_document(tables, path, caption_style=None, append=False):
    """Write one or more tables to a Word document.

    Args:
        tables: A single :class:`Table` or an iterable of them.
        path: Destination ``.docx`` path.
        caption_style: Optional named Word style for captions. When the
            style is absent from the template the caption is written as
            bold body text instead.
        append: Add the tables to the end of an existing document at
            ``path`` rather than starting a new one. Lets the per-task
            scripts build up a single file of numbered tables.

    Returns:
        The path written.
    """
    if isinstance(tables, Table):
        tables = [tables]
    tables = list(tables)

    reopened = append and os.path.exists(path)
    document = Document(path) if reopened else Document()

    for index, table in enumerate(tables):
        if index or reopened:
            document.add_paragraph()

        caption = document.add_paragraph()
        if caption_style:
            try:
                caption.style = document.styles[caption_style]
            except KeyError:
                pass
        _add_runs(caption, f'<b>Table {table.number}.</b> {table.caption}')

        _append_table(document, table)

        for letter, note in zip(FOOTNOTE_LETTERS, table.footnotes):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            marker = paragraph.add_run(letter)
            marker.font.superscript = True
            _add_runs(paragraph, note)

    document.save(path)
    return path


def _append_table(document, table):
    """Append one :class:`Table` to a Word document.

    Args:
        document: The ``docx`` document.
        table: The table to append.

    Returns:
        The created ``docx`` table.
    """
    nested = table.is_nested
    n_stub = 2 if nested else 1
    n_cols = n_stub + len(table.columns)

    grid = document.add_table(rows=1, cols=n_cols)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = True
    try:
        grid.style = document.styles['Table Grid']
    except KeyError:
        pass

    header = grid.rows[0]
    _repeat_header_row(header)

    if nested:
        stub = header.cells[0].merge(header.cells[1])
    else:
        stub = header.cells[0]
    _write_cell(stub, table.stub_head, bold=True)

    for offset, heading in enumerate(table.columns):
        _write_cell(header.cells[n_stub + offset], heading, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in table.rows:
        cells = grid.add_row().cells

        if row.kind == 'category':
            # JMIR: merge the category heading with the subcategory
            # column so the subcategories below it appear indented.
            merged = cells[0].merge(cells[1]) if nested else cells[0]
            _write_cell(merged, row.label, bold=True)
            continue

        if nested and row.level == 0:
            merged = cells[0].merge(cells[1])
            _write_cell(merged, row.label)
        elif nested:
            _write_cell(cells[0], '')
            _write_cell(cells[1], row.label)
        else:
            _write_cell(cells[0], row.label)

        for offset, value in enumerate(row.cells):
            _write_cell(cells[n_stub + offset], value,
                        align=WD_ALIGN_PARAGRAPH.CENTER)

    return grid
